#!/usr/bin/env python3
#
# freemind_to_word.py
#
# Copyright LiKneu 2019
#

import lxml.etree as ET # handling XML
import docx             # handling of MS Word documents 

def main(args):
    return 0

def to_word(input_file, output_file):
    '''Converts the Freeemind XML file into an MS Mord DOCX file.'''

    print('Converting to Word')
    print('Input file : ' + input_file)
    print('Output file: ' + output_file)
    
    # Prepare the Fremmind XML tree
    # Read freemind XML file into variable
    fm_tree = ET.parse(input_file)
    # Get the root element of the Freemind XML tree
    fm_root = fm_tree.getroot()
    
    doc = docx.Document()
    
    # Get the title of the freemind root node and use it as title of the
    # cover page
    title_cover = fm_root.xpath('/map/node/@TEXT')[0]
    main_title = doc.add_heading(title_cover, 0)
    
    # Add a paragraph below the title heading. I found no other way to directly
    # add a page break after the heading.
    para = doc.add_paragraph(' ')
    doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    
    # Walk through all <node> tags of the freemind file
    for fm_node in fm_root.iter('node'):        
        # Get the content of the "TEXT" attribute of the node
        wd_title = fm_node.get("TEXT")
        # Calculate how deep we have gone into the tree structure to choose a
        # fitting heading style
        wd_title_level = fm_tree.getpath(fm_node).count('/')-1
        # The 1st 4 levels get different heading styles..
        if wd_title_level <= 4:
            doc.add_heading(wd_title, wd_title_level)
        else:
            #..all the other levels stick with one style
            doc.add_paragraph(wd_title, 'Heading5')
        
        # Check if node has an attached note <richcontent>
        fm_note = fm_node.xpath('normalize-space(./richcontent)')
        # If yes, remove all html tags and store the remaining text in a
        # Project <Note> tag
        if fm_note:
            doc.add_paragraph(fm_note, 'Normal')

    # Write the Word DOCX file to disc
    doc.save(output_file)

if __name__ == '__main__':
    import sys
    sys.exit(main(sys.argv))
